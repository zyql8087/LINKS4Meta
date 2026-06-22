"""Generate V4 experiment report as a Word (.docx) document.
V2: fixes S0 table, adds ready_for_rl, B1 top-k/MRR, F score comparison, multi-seed.
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parents[1]
OUTPUT_DIR = PROJECT_ROOT / "demo" / "outputs"
REPORT_PATH = OUTPUT_DIR / "V4_experiment_report_v3.docx"

FAMILIES = ["6bar", "7bar", "8bar", "9bar"]

# ── data paths ──────────────────────────────────────────────────────────
DATA_FILES = {
    "s0": OUTPUT_DIR / "forward_precheck_v4.json",
    "a_baseline": OUTPUT_DIR / "readout_v4_A_baseline_s23.json",
    "a_baseline_s89": OUTPUT_DIR / "readout_v4_A_baseline_s89.json",
    "a_schemed": OUTPUT_DIR / "readout_v4_A_schemed_s23.json",
    "a_seed47": OUTPUT_DIR / "readout_v4_A_seed47.json",
    "a_seed89": OUTPUT_DIR / "readout_v4_A_seed89.json",
    "b1": OUTPUT_DIR / "readout_v4_B_ranking_quality.json",
    "c_9bar": OUTPUT_DIR / "readout_v4_C_9bar.json",
    "d_noprior": OUTPUT_DIR / "readout_v4_D_noprior.json",
    "d_highcap": OUTPUT_DIR / "readout_v4_D_highcap.json",
    "d_topk1": OUTPUT_DIR / "readout_v4_D_topk1.json",
    "f_shuffle": OUTPUT_DIR / "readout_v4_F_shuffle.json",
    "f_shuffle_seed47": OUTPUT_DIR / "readout_v4_F_shuffle_seed47.json",
    "f_shuffle_seed89": OUTPUT_DIR / "readout_v4_F_shuffle_seed89.json",
}


def _load(key: str) -> dict | None:
    p = DATA_FILES[key]
    if not p.exists():
        return None
    with p.open("r", encoding="utf-8") as f:
        return json.load(f)


# ── helpers ─────────────────────────────────────────────────────────────

def _set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {qn("w:fill"): color_hex, qn("w:val"): "clear"})
    shading.append(shd)


def add_table(doc, headers: list[str], rows: list[list], *, header_color="2E75B6"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(cell, header_color)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                _set_cell_shading(cell, "F2F2F2")
    return table


def _fmt(v, decimals=4):
    if v is None:
        return "—"
    if isinstance(v, float):
        return f"{v:.{decimals}f}"
    return str(v)


def _pct(v):
    if v is None:
        return "—"
    return f"{v * 100:.2f}%"


def _raw_joint_score(sd: dict) -> tuple[float | None, bool]:
    """Get raw joint_score: prefer explicit field, else estimate from assignment_score.

    Returns (value, is_estimate). is_estimate=True when computed from assignment_score.
    """
    rj = sd.get("mean_raw_joint_score")
    if rj is not None:
        return float(rj), False
    # fallback: assignment_score = -joint_score + prior, estimate prior ≈ 0.047
    # This is approximate since structural score varies per sample.
    assignment = sd.get("mean_assignment_score")
    if assignment is not None:
        return float(-assignment - 0.047), True
    return None, True


def _add_bold_text(paragraph, bold_text, normal_text):
    run = paragraph.add_run(bold_text)
    run.bold = True
    paragraph.add_run(normal_text)


# ── section builders ────────────────────────────────────────────────────

def build_title(doc):
    doc.add_heading("GraphMetaMat-LINKS V4 实验报告", level=0)
    doc.add_paragraph("SurrogateTargetReadoutAssignment (scheme_d) 验证实验")
    doc.add_paragraph("")


# ─────────────────────────────────────────────────────────────────────
# S0
# ─────────────────────────────────────────────────────────────────────

def build_s0(doc):
    doc.add_heading("S0 — Forward Surrogate 前检", level=1)
    doc.add_paragraph(
        "评估 forward surrogate 在各 mechanism family 上的预测质量。"
        "若 9bar 显著劣于其他 family，则下游 readout 精度将受制约。"
    )
    data = _load("s0")
    if data is None:
        doc.add_paragraph("[数据文件缺失]")
        return

    # S0 data is under current_model.test.per_family
    cm = data.get("current_model", {})
    test = cm.get("test", {})
    pf = test.get("per_family", {})
    overall = test.get("overall", {})
    gate = data.get("gate", {})

    # main metrics table
    headers = ["Family", "foot_path_error", "knee_nmae", "ankle_nmae", "foot_chamfer", "ankle_std_ratio"]
    rows = []
    for fam in FAMILIES:
        d = pf.get(fam, {})
        rows.append([
            fam,
            _fmt(d.get("foot_path_error"), 4),
            _fmt(d.get("knee_nmae"), 4),
            _fmt(d.get("ankle_nmae"), 4),
            _fmt(d.get("foot_chamfer"), 4),
            _fmt(d.get("ankle_std_ratio"), 3),
        ])
    # overall row
    rows.append([
        "Overall",
        _fmt(overall.get("foot_path_error"), 4),
        _fmt(overall.get("knee_nmae"), 4),
        _fmt(overall.get("ankle_nmae"), 4),
        _fmt(overall.get("foot_chamfer"), 4),
        _fmt(overall.get("ankle_std_ratio"), 3),
    ])
    add_table(doc, headers, rows)

    # gate status
    doc.add_paragraph("")
    doc.add_heading("Gate Status", level=2)
    headers_g = ["Gate", "Value"]
    rows_g = [
        ["semantic_collapse", str(gate.get("semantic_collapse", "—"))],
        ["ready_for_rl", str(gate.get("ready_for_rl", "—"))],
        ["bar89_out_of_control", str(gate.get("bar89_out_of_control", "—"))],
        ["family_imbalance", str(gate.get("family_imbalance", "—"))],
        ["stronger_than_baselines (retrieval)", str(gate.get("stronger_than_baselines", {}).get("retrieval", "—"))],
    ]
    add_table(doc, headers_g, rows_g, header_color="C0504D")

    # semantic ablation
    sa = data.get("semantic_ablation", {})
    if sa:
        doc.add_paragraph("")
        doc.add_heading("Semantic Ablation Degradation Ratios", level=2)
        headers_sa = ["Metric", "Degradation Ratio"]
        rows_sa = [[k.replace("_degradation_ratio", "").replace("_", " "), _fmt(v, 2)] for k, v in sa.items()]
        add_table(doc, headers_sa, rows_sa, header_color="9BBB59")

    doc.add_paragraph("")
    p = doc.add_paragraph()
    _add_bold_text(p, "结论: ",
        "9bar ankle_nmae = 0.859（2.19× 劣于 6bar 的 0.393）。"
        "9bar 的 readout 上限被 forward surrogate 限制。"
    )
    p2 = doc.add_paragraph()
    _add_bold_text(p2, "ready_for_rl=False 的含义: ",
        "这说明当前 surrogate 虽然能支持 readout 排序（V4 实验已证明），"
        "但还没有通过自己设定的 RL/端到端门槛。"
        "论文中应明确表述：V4 证明的是 readout 选择，不是完整 RL rollout 已可靠。"
    )


# ─────────────────────────────────────────────────────────────────────
# A
# ─────────────────────────────────────────────────────────────────────

def _scheme_d_overall_acc(data_d, mode="graph_target"):
    sd = data_d.get("eval", {}).get(mode, {}).get("scheme_d", {})
    return sd.get("exact_chain_accuracy")


def _scheme_d_family_accs(data_d, mode="graph_target"):
    sd = data_d.get("eval", {}).get(mode, {}).get("scheme_d", {})
    by_fam = sd.get("by_family", {})
    return {f: by_fam.get(f, {}).get("exact_chain_accuracy") for f in FAMILIES}


def build_a(doc):
    doc.add_heading("A — Full-split 主实验", level=1)

    # baseline
    doc.add_heading("A1: Baseline（scheme_a + scheme_c）", level=2)
    data_bl = _load("a_baseline")
    if data_bl is None:
        doc.add_paragraph("[数据文件缺失]")
    else:
        eval_results = data_bl.get("eval", {})
        headers = ["Scheme", "Eval Mode", "Exact Chain Acc", "Knee Acc", "Ankle Acc", "Foot Acc"]
        rows = []
        for mode_name in ["graph_target", "graph_only", "graph_motion_target"]:
            mode_data = eval_results.get(mode_name, {})
            for scheme_name in ["scheme_a", "scheme_c"]:
                sd = mode_data.get(scheme_name, {})
                if not sd:
                    continue
                rows.append([
                    scheme_name, mode_name,
                    _pct(sd.get("exact_chain_accuracy")),
                    _pct(sd.get("knee_accuracy")),
                    _pct(sd.get("ankle_accuracy")),
                    _pct(sd.get("foot_accuracy")),
                ])
        add_table(doc, headers, rows)
        doc.add_paragraph("")
        p = doc.add_paragraph()
        _add_bold_text(p, "关键发现: ",
            "scheme_a 在 deployment mode (graph_target) 下完全失败 (0.0%)。"
            "scheme_c graph_only (71.6%) ≥ graph_target (70.3%)，"
            "表明 SlotPointer 未能有效利用 target 曲线信息。"
        )

        # scheme_c multi-seed
        def _sc_acc(data_dict, mode="graph_target"):
            if data_dict is None:
                return None
            return data_dict.get("eval", {}).get(mode, {}).get("scheme_c", {}).get("exact_chain_accuracy")

        sc_seeds = [("seed 23", _sc_acc(data_bl))]
        for key, label in [("a_seed47", "seed 47"), ("a_baseline_s89", "seed 89")]:
            d = _load(key)
            acc = _sc_acc(d)
            if acc is not None:
                sc_seeds.append((label, acc))

        if len(sc_seeds) > 1:
            doc.add_paragraph("")
            doc.add_heading("Scheme_c Multi-seed Stability (graph_target)", level=3)
            headers_ms = ["Seed", "Scheme_c Acc"]
            rows_ms = [[label, _pct(acc)] for label, acc in sc_seeds]
            if len(sc_seeds) > 2:
                import numpy as np
                accs = [x[1] for x in sc_seeds]
                rows_ms.append(["mean ± std", f"{np.mean(accs)*100:.2f}% ± {np.std(accs)*100:.2f}%"])
            add_table(doc, headers_ms, rows_ms)
            doc.add_paragraph("")
            p_ms = doc.add_paragraph()
            _add_bold_text(p_ms, "注意: ",
                "scheme_c 精度在不同 seed 下波动较大（训练依赖随机初始化和数据采样），"
                "而 scheme_d 是确定性的（0 方差）。两者的多 seed 比较意义不同。"
            )

    # scheme_d
    doc.add_heading("A2: Scheme_d（SurrogateTarget）", level=2)
    data_d = _load("a_schemed")
    if data_d is None:
        doc.add_paragraph("[数据文件缺失]")
    else:
        eval_results = data_d.get("eval", {})
        headers = ["Eval Mode", "Exact Chain Acc", "Knee Acc", "Ankle Acc", "Foot Acc", "Avg Time (s)"]
        rows = []
        for mode_name in ["graph_target", "graph_only", "graph_motion_target"]:
            sd = eval_results.get(mode_name, {}).get("scheme_d", {})
            if not sd:
                continue
            rows.append([
                mode_name,
                _pct(sd.get("exact_chain_accuracy")),
                _pct(sd.get("knee_accuracy")),
                _pct(sd.get("ankle_accuracy")),
                _pct(sd.get("foot_accuracy")),
                _fmt(sd.get("mean_assignment_time_sec"), 3),
            ])
        add_table(doc, headers, rows)

        # per-family
        doc.add_paragraph("")
        doc.add_heading("Per-family (graph_target mode)", level=3)
        gt = eval_results.get("graph_target", {}).get("scheme_d", {})
        by_fam = gt.get("by_family", {})
        headers2 = ["Family", "Exact Chain Acc", "Knee Acc", "Ankle Acc", "Foot Acc", "Candidates", "Cap Rate"]
        rows2 = []
        for fam in FAMILIES:
            fd = by_fam.get(fam, {})
            cc = fd.get("candidate_count", {})
            rows2.append([
                fam,
                _pct(fd.get("exact_chain_accuracy")),
                _pct(fd.get("knee_accuracy")),
                _pct(fd.get("ankle_accuracy")),
                _pct(fd.get("foot_accuracy")),
                _fmt(cc.get("mean"), 0),
                _pct(cc.get("reaches_cap_rate")),
            ])
        add_table(doc, headers2, rows2)

        # multi-seed if available
        multi_seed_data = []
        for seed_label, seed_key in [("seed 47", "a_seed47"), ("seed 89", "a_seed89")]:
            sd_data = _load(seed_key)
            if sd_data is not None:
                acc = _scheme_d_overall_acc(sd_data)
                fam_accs = _scheme_d_family_accs(sd_data)
                multi_seed_data.append((seed_label, acc, fam_accs))

        if multi_seed_data:
            doc.add_paragraph("")
            doc.add_heading("Multi-seed Stability (scheme_d graph_target)", level=3)
            # collect all seeds
            seed_23_acc = _scheme_d_overall_acc(data_d)
            seed_23_fams = _scheme_d_family_accs(data_d)
            all_seeds = [("seed 23", seed_23_acc, seed_23_fams)] + multi_seed_data
            headers_ms = ["Seed", "Overall"] + FAMILIES
            rows_ms = []
            for s_label, s_acc, s_fams in all_seeds:
                rows_ms.append([s_label, _pct(s_acc)] + [_pct(s_fams.get(f)) for f in FAMILIES])
            # mean/std
            import numpy as np
            all_overall = [x[1] for x in all_seeds if x[1] is not None]
            all_fams = {f: [x[2].get(f) for x in all_seeds if x[2].get(f) is not None] for f in FAMILIES}
            if len(all_overall) > 1:
                rows_ms.append([
                    "mean ± std",
                    f"{np.mean(all_overall)*100:.2f}% ± {np.std(all_overall)*100:.2f}%",
                ] + [
                    f"{np.mean(all_fams[f])*100:.2f}% ± {np.std(all_fams[f])*100:.2f}%" if all_fams[f] else "—"
                    for f in FAMILIES
                ])
            add_table(doc, headers_ms, rows_ms)

        doc.add_paragraph("")
        p = doc.add_paragraph()
        _add_bold_text(p, "结论: ",
            "scheme_d 在 graph_target mode 下 exact_chain_accuracy = 87.85%，"
            "显著优于 scheme_c 的 ~68.6%（三 seed mean, +19.2%）。"
            "SurrogateTarget 在纯部署模式（graph+target，无训练数据）下大幅领先 SlotPointer。"
        )


# ─────────────────────────────────────────────────────────────────────
# B1
# ─────────────────────────────────────────────────────────────────────

def build_b1(doc):
    doc.add_heading("B1 — Surrogate Ranking Quality", level=1)
    doc.add_paragraph(
        "评估 forward surrogate 对 candidate leg chains 的排序能力。"
        "使用 exact match（二值）作为 oracle score，对比 surrogate 预测排序。"
        "joint_score 越低越好（误差），排序时取负后按降序排列。"
    )
    data = _load("b1")
    if data is None:
        doc.add_paragraph("[数据文件缺失]")
        return
    overall = data.get("overall", {})
    pf = data.get("per_family", {})

    # primary table: top-k / MRR / rank / NDCG
    headers = ["Family", "Top-1", "Top-3", "Top-5", "MRR", "NDCG@1",
               "Median Truth Rank", "Mean Truth Rank", "Candidates", "Coverage"]
    rows = []

    def _row_from_dict(d, label):
        return [
            label,
            _pct(d.get("top_1")),
            _pct(d.get("top_3")),
            _pct(d.get("top_5")),
            _fmt(d.get("mrr"), 3) if d.get("mrr") is not None else "—",
            _fmt(d.get("mean_ndcg_at_1"), 3),
            _fmt(d.get("median_truth_rank"), 0),
            _fmt(d.get("mean_truth_rank"), 1),
            _fmt(d.get("mean_candidate_count"), 0),
            _pct(d.get("truth_candidate_coverage")),
        ]

    rows.append(_row_from_dict(overall, "Overall"))
    for fam in FAMILIES:
        rows.append(_row_from_dict(pf.get(fam, {}), fam))
    add_table(doc, headers, rows)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    _add_bold_text(p, "结论: ",
        "Surrogate ranking 质量随 family 复杂度递减。"
        "Overall Top-1 = 88.26%，与 NDCG@1 = 0.883 数值一致，验证排序正确。"
        "9bar Top-1 = 65.98%（median truth rank = 1，mean = 7.81），"
        "说明正确候选多数时候仍排第一，但有 ~34% 的样本被其他候选超越。"
    )


# ─────────────────────────────────────────────────────────────────────
# C
# ─────────────────────────────────────────────────────────────────────

def build_c(doc):
    doc.add_heading("C — 9bar 根因诊断", level=1)
    doc.add_paragraph("只用 9bar 样本（1120 条），train_per_family=128，对比各 scheme 在 9bar 上的表现。")
    data = _load("c_9bar")
    if data is None:
        doc.add_paragraph("[数据文件缺失]")
        return
    eval_results = data.get("eval", {})

    headers = ["Scheme", "Eval Mode", "Exact Chain Acc", "Knee Acc", "Ankle Acc", "Foot Acc"]
    rows = []
    for mode_name in ["graph_target", "graph_motion_target"]:
        mode_data = eval_results.get(mode_name, {})
        for scheme_name in ["scheme_a", "scheme_c", "scheme_d"]:
            sd = mode_data.get(scheme_name, {})
            if not sd:
                continue
            by_fam = sd.get("by_family", {}).get("9bar", {})
            rows.append([
                scheme_name, mode_name,
                _pct(by_fam.get("exact_chain_accuracy")),
                _pct(by_fam.get("knee_accuracy")),
                _pct(by_fam.get("ankle_accuracy")),
                _pct(by_fam.get("foot_accuracy")),
            ])
    add_table(doc, headers, rows)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    _add_bold_text(p, "关键发现: ",
        "在纯 deployment mode (graph+target) 下，scheme_d (67.95%) > scheme_c (51.79%)，差距 +16.2%。"
        "scheme_c 在有 motion data 时接近完美 (99.38%)，"
        "说明 SlotPointer 能学到 motion pattern，但不能从 target curve 推断。"
    )


# ─────────────────────────────────────────────────────────────────────
# D
# ─────────────────────────────────────────────────────────────────────

def build_d(doc):
    doc.add_heading("D — 参数消融（scheme_d, graph_target）", level=1)
    doc.add_paragraph("验证 structural_prior_weight, max_candidates_cap, top_k 三个超参数对精度的影响。")

    def _family_accs(key):
        data = _load(key)
        if data is None:
            return None, None
        sd = data.get("eval", {}).get("graph_target", {}).get("scheme_d", {})
        by_fam = sd.get("by_family", {})
        overall_acc = sd.get("exact_chain_accuracy")
        fam_accs = {f: by_fam.get(f, {}).get("exact_chain_accuracy") for f in FAMILIES}
        return overall_acc, fam_accs

    def_acc, def_fams = _family_accs("a_schemed")
    np_acc, np_fams = _family_accs("d_noprior")
    hc_acc, hc_fams = _family_accs("d_highcap")
    tk_acc, tk_fams = _family_accs("d_topk1")

    headers = ["Variant", "Overall", "6bar", "7bar", "8bar", "9bar"]
    rows = []
    for label, ov, fams in [
        ("Default (prior=0.05, cap=256, top_k=5)", def_acc, def_fams),
        ("D-no-prior (prior=0)", np_acc, np_fams),
        ("D-high-cap (cap=512)", hc_acc, hc_fams),
        ("D-topk1 (top_k=1)", tk_acc, tk_fams),
    ]:
        if fams is None:
            rows.append([label] + ["N/A"] * 5)
            continue
        rows.append([label, _pct(ov)] + [_pct(fams.get(f)) for f in FAMILIES])
    add_table(doc, headers, rows)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    _add_bold_text(p, "结论: ",
        "D-no-prior 将 structural_prior_weight 设为 0 后，"
        "overall 从 87.85% 提升至 ~88.64%（+0.79%）。"
        "差异不算巨大，但不应简单说“完全无影响”。"
        "建议后续把默认 structural_prior_weight 改为 0 或至少在多 seed 下确认。"
        "D-high-cap 和 D-topk1 与 default 完全一致。"
        "瓶颈在 forward surrogate 的 ranking 质量，不在决策超参。"
    )


# ─────────────────────────────────────────────────────────────────────
# F
# ─────────────────────────────────────────────────────────────────────

def build_f(doc):
    doc.add_heading("F — Target Shuffle 反证", level=1)
    doc.add_paragraph(
        "将 target curve 在同 family 内随机打乱后评估。"
        "若 accuracy 跌至接近随机水平，说明 scheme_d 确实依赖 target curve 信息。"
    )
    data_real = _load("a_schemed")
    data_shuf = _load("f_shuffle")
    if data_real is None or data_shuf is None:
        doc.add_paragraph("[数据文件缺失]")
        return

    real_sd = data_real.get("eval", {}).get("graph_target", {}).get("scheme_d", {})
    shuf_sd = data_shuf.get("eval", {}).get("graph_target", {}).get("scheme_d", {})
    real_by_fam = real_sd.get("by_family", {})
    shuf_by_fam = shuf_sd.get("by_family", {})

    # accuracy comparison
    headers = ["Family", "Real Target Acc", "Shuffled Acc", "Drop"]
    rows = []
    for fam in FAMILIES:
        r = real_by_fam.get(fam, {}).get("exact_chain_accuracy")
        s = shuf_by_fam.get(fam, {}).get("exact_chain_accuracy")
        drop = (s - r) if (r is not None and s is not None) else None
        rows.append([fam, _pct(r), _pct(s), _pct(drop)])
    # overall
    r_ov = real_sd.get("exact_chain_accuracy")
    s_ov = shuf_sd.get("exact_chain_accuracy")
    drop_ov = (s_ov - r_ov) if (r_ov is not None and s_ov is not None) else None
    rows.append(["Overall", _pct(r_ov), _pct(s_ov), _pct(drop_ov)])
    add_table(doc, headers, rows)

    # raw score comparison
    doc.add_paragraph("")
    doc.add_heading("Surrogate Score Comparison", level=2)
    doc.add_paragraph(
        "mean_assignment_score = mean(-joint_score + prior)，越高越好。"
        "mean_raw_joint_score = mean(joint_score)，越低越好（误差度量）。"
        "两个指标方向相反但含义一致：real target 匹配更好。"
    )
    real_score = real_sd.get("mean_assignment_score")
    shuf_score = shuf_sd.get("mean_assignment_score")
    real_raw, real_est = _raw_joint_score(real_sd)
    shuf_raw, shuf_est = _raw_joint_score(shuf_sd)
    def _fmt_raw(v, is_est):
        if v is None:
            return "—"
        prefix = "~" if is_est else ""
        return f"{prefix}{v:.4f}"
    headers_s = ["Condition", "Assignment Score (↑=good)", "Raw Joint Score (↓=good)", "Exact Chain Acc"]
    rows_s = [
        ["Real target", _fmt(real_score, 4), _fmt_raw(real_raw, real_est), _pct(r_ov)],
        ["Shuffled target", _fmt(shuf_score, 4), _fmt_raw(shuf_raw, shuf_est), _pct(s_ov)],
    ]
    add_table(doc, headers_s, rows_s, header_color="C0504D")

    # multi-seed shuffle if available
    seed_results = [("seed 23", s_ov)]
    for seed_key, seed_label in [("f_shuffle_seed47", "seed 47"), ("f_shuffle_seed89", "seed 89")]:
        sd = _load(seed_key)
        if sd is not None:
            s = sd.get("eval", {}).get("graph_target", {}).get("scheme_d", {}).get("exact_chain_accuracy")
            if s is not None:
                seed_results.append((seed_label, s))

    if len(seed_results) > 1:
        doc.add_paragraph("")
        doc.add_heading("Multi-seed Shuffle Stability", level=3)
        headers_ms = ["Seed", "Shuffled Acc"]
        rows_ms = [[label, _pct(acc)] for label, acc in seed_results]
        if len(seed_results) > 2:
            import numpy as np
            accs = [x[1] for x in seed_results]
            rows_ms.append(["mean ± std", f"{np.mean(accs)*100:.2f}% ± {np.std(accs)*100:.2f}%"])
        add_table(doc, headers_ms, rows_ms)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    raw_ratio = ""
    if real_raw and shuf_raw and real_raw > 0:
        est_note = "（估算）" if (real_est or shuf_est) else ""
        raw_ratio = f"（raw joint_score 从 {real_raw:.4f} 恶化至 {shuf_raw:.4f}，约 {shuf_raw/real_raw:.1f}×{est_note}）"
    _add_bold_text(p, "结论: ",
        f"Shuffle 后 accuracy 从 {_pct(r_ov)} 暴跌至 {_pct(s_ov)}。"
        f"{raw_ratio}"
        "assignment_score 下降和 raw joint_score 上升方向一致，"
        "确认 surrogate 对 target 扰动敏感，accuracy 下降是 score 恶化的直接后果。"
    )


# ─────────────────────────────────────────────────────────────────────
# Limitations
# ─────────────────────────────────────────────────────────────────────

def build_limitations(doc):
    doc.add_heading("局限性与未来工作", level=1)

    doc.add_heading("不可执行的实验", level=2)
    headers = ["实验", "原因"]
    rows = [
        ["E (Phase5 diagnostics)", "需要 RL checkpoint (model_inverse_rl.pt)，不存在"],
        ["H (MCTS)", "需要 IL + RL checkpoint，不存在"],
        ["G (physical feasibility)", "需要单髋驱动仿真器，未实现"],
    ]
    add_table(doc, headers, rows)

    doc.add_paragraph("")
    doc.add_heading("当前结论边界", level=2)
    p = doc.add_paragraph()
    _add_bold_text(p, "可以写: ",
        "SurrogateTargetReadoutAssignment substantially improves deployment-time semantic readout "
        "under graph+target input, achieving 87.85% exact-chain accuracy over the full test split "
        "and outperforming SlotPointer by 17.6 percentage points. The target-shuffle control confirms "
        "strong target dependence. However, performance on 9bar remains constrained by forward surrogate "
        "ranking quality, and end-to-end RL/physical feasibility validation remains future work."
    )
    doc.add_paragraph("")
    p2 = doc.add_paragraph()
    _add_bold_text(p2, "不要写: ",
        "“The full inverse generation pipeline is validated.” "
        "“9bar failure is fully solved.”"
    )

    doc.add_paragraph("")
    doc.add_heading("已完成的改进", level=2)
    steps = [
        "[已完成] S0 表格修复，正式列出 6bar/7bar/8bar/9bar 的 forward metrics。",
        "[已完成] seed 23/47/89 多 seed 实验，给 A/F 结果加 mean/std。",
        "[已完成] F shuffle 增加 raw surrogate_joint_score 对比。",
        "[已完成] B1 修复排序方向 bug，Top-1/NDCG@1 现在一致。",
        "[已完成] 将 structural_prior_weight=0 纳入默认候选（+0.79%）。",
        "[待补] E/H 需要 RL checkpoint，G 需要仿真器——"
        "若论文目标包含完整系统，必须补 E/H；"
        "若只聚焦 readout，则明确把 E/H/G 放入 limitation/future work。",
    ]
    for s in steps:
        doc.add_paragraph(s, style="List Bullet")


# ─────────────────────────────────────────────────────────────────────
# Conclusion
# ─────────────────────────────────────────────────────────────────────

def build_conclusion(doc):
    doc.add_heading("总体结论", level=1)
    conclusions = [
        "SurrogateTarget readout (scheme_d) 在纯 deployment mode 下显著优于 SlotPointer (scheme_c): "
        "87.85% vs ~68.6%（三 seed mean, +19.2%）。scheme_d 是确定性的（三 seed 完全一致），"
        "scheme_c 在不同 seed 下波动较大（seed23=70.3%, seed47=62.8%, seed89=72.8%, std=±4.2%）。",
        "SlotPointer 的 target 利用能力弱: graph_only (71.6%) ≥ graph_target (70.3%)，"
        "且 9bar graph_target 仅 51.79%。",
        "9bar 性能瓶颈在 forward surrogate 排序质量: S0 显示 ankle_nmae=0.859（2.19× 劣于 6bar），"
        "B1 显示 Top-1=65.98%（vs 6bar 的 96.36%），直接制约 readout accuracy。",
        "ready_for_rl=False 说明当前结果仅证明 readout 选择能力，不是完整 RL rollout 已可靠。",
        "决策超参影响极小: D 实验表明 prior=0 反而轻微提升 (+0.79%)，"
        "cap 和 top_k 无影响。建议将默认 structural_prior_weight 改为 0。",
        "Target 信息是有效信号: F 实验 shuffle 后 accuracy 从 87.85% 暴跌至 ~12%，"
        "确认 surrogate 依赖 target curve 完成 readout。",
    ]
    for i, c in enumerate(conclusions, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(c)


# ─────────────────────────────────────────────────────────────────────
# Appendix
# ─────────────────────────────────────────────────────────────────────

def build_appendix(doc):
    doc.add_heading("附录: 输出文件清单", level=1)
    headers = ["实验", "输出文件"]
    rows = [
        ["S0", "forward_precheck_v4.json"],
        ["A baseline (seed 23)", "readout_v4_A_baseline_s23.json"],
        ["A baseline (seed 89)", "readout_v4_A_baseline_s89.json"],
        ["A scheme_d (seed 23)", "readout_v4_A_schemed_s23.json"],
        ["A scheme_d (seed 47)", "readout_v4_A_seed47.json"],
        ["A scheme_d (seed 89)", "readout_v4_A_seed89.json"],
        ["B1", "readout_v4_B_ranking_quality.json"],
        ["C 9bar", "readout_v4_C_9bar.json"],
        ["D no-prior", "readout_v4_D_noprior.json"],
        ["D high-cap", "readout_v4_D_highcap.json"],
        ["D topk1", "readout_v4_D_topk1.json"],
        ["F shuffle (seed 23)", "readout_v4_F_shuffle.json"],
        ["F shuffle (seed 47)", "readout_v4_F_shuffle_seed47.json"],
        ["F shuffle (seed 89)", "readout_v4_F_shuffle_seed89.json"],
    ]
    add_table(doc, headers, rows)
    doc.add_paragraph("")
    doc.add_paragraph("所有文件位于: demo/outputs/ 目录下。")


# ── main ────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    build_title(doc)
    build_s0(doc)
    build_a(doc)
    build_b1(doc)
    build_c(doc)
    build_d(doc)
    build_f(doc)
    build_limitations(doc)
    build_conclusion(doc)
    build_appendix(doc)

    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(REPORT_PATH))
    print(f"Report saved to: {REPORT_PATH}")


if __name__ == "__main__":
    main()
